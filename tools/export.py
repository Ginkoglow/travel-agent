import os
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from docx import Document

def export_plan_to_pdf(content: str, filename: str = None) -> str:
    """
    导出旅行计划为 PDF，返回文件绝对路径。
    """
    if filename is None:
        filename = f"travel_plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    
    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4

    # 标题
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, "旅行智能助手 - 专属攻略")
    c.line(50, height - 55, width - 50, height - 55)

    # 正文
    c.setFont("Helvetica", 11)
    y = height - 80
    lines = content.split('\n')
    for line in lines:
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 50
        # 自动换行
        wrapped = simpleSplit(line, "Helvetica", 11, width - 100)
        for wline in wrapped:
            c.drawString(50, y, wline)
            y -= 16
        y -= 5

    c.save()
    return os.path.abspath(filename)

def export_to_word(content: str, filename: str = "旅行攻略.docx") -> str:
    doc = Document()
    doc.add_heading('旅行智能助手 - 专属攻略', 0)
    doc.add_paragraph(content)
    doc.save(filename)
    return os.path.abspath(filename)